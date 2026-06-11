#!/usr/bin/env python3
"""
Bulk generator for Slovak "Zálohová faktúra" DOCX/PDF files plus sidecar manifests.

The script produces:
- DOCX invoices
- optional text-based PDF invoices generated directly from Python
- invoice/customer CSV and JSON manifests for downstream reconciliation tests

Note: the QR payload below is plain payment/invoice text for testing. If you need
bank-app-compatible Slovak PAY by square QR codes, replace make_qr_payload() with
proper PAY by square encoding.
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import json
import random
import re
from dataclasses import dataclass, replace
from decimal import Decimal, ROUND_HALF_UP
from functools import lru_cache
from pathlib import Path
from typing import Any, Iterable, Optional
from xml.sax.saxutils import escape

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.table import Table
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table as PdfTable, TableStyle


# -----------------------------
# Data models
# -----------------------------


@dataclass(frozen=True)
class Party:
    name: str
    street: str
    city_country: str
    ico: str = ""
    dic: str = ""
    ic_dph: str = ""
    register: str = ""
    vat_payer: str = "Nie"
    phone: str = ""
    email: str = ""
    web: str = ""
    iban: str = ""
    swift: str = ""
    bank_account: str = ""
    contact: str = ""
    note: str = ""
    customer_id: str = ""
    bank_code: str = ""
    account_number: str = ""
    bank_name: str = ""
    monthly_fee: str = ""
    annual_extra_fee: str = ""
    annual_extra_interval_months: str = ""
    payment_method: str = ""
    vat_status: str = ""
    status: str = ""

    @property
    def iban_compact(self) -> str:
        return compact_iban(self.iban)


@dataclass(frozen=True)
class InvoiceData:
    invoice_no: str
    variable_symbol: str
    issue_date: dt.date
    due_date: dt.date
    supplier: Party
    customer: Party
    item_description: str
    quantity: Decimal
    unit: str
    unit_price: Decimal
    prepared_by: str
    received_by: str = ""
    note: str = "Tento dokument je vzorovo vyplnený testovacími údajmi."

    @property
    def total(self) -> Decimal:
        return money(self.quantity * self.unit_price)


@dataclass(frozen=True)
class PlannedInvoice:
    batch_id: str
    invoice: InvoiceData
    charge_type: str
    amount_bucket: str
    payment_scenario: str
    simulated_paid_total: Decimal
    simulated_split_amounts: tuple[Decimal, ...]
    reference_text_template: str
    filename_base: str

    @property
    def docx_filename(self) -> str:
        return f"{self.filename_base}.docx"

    @property
    def pdf_filename(self) -> str:
        return f"{self.filename_base}.pdf"


# -----------------------------
# Formatting helpers
# -----------------------------


def money(value: Decimal | float | int | str) -> Decimal:
    return Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def money_eur(value: Decimal | float | int | str) -> str:
    value = money(value)
    s = f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", " ")
    return f"{s} €"


def money_number(value: Decimal | float | int | str) -> str:
    return money_eur(value).replace(" €", "")


def decimal_string(value: Decimal | float | int | str) -> str:
    return f"{money(value):.2f}"


def compact_decimal_string(value: Decimal) -> str:
    normalized = value.normalize()
    return format(normalized, "f")


ONES = {
    0: "nula",
    1: "jeden",
    2: "dva",
    3: "tri",
    4: "štyri",
    5: "päť",
    6: "šesť",
    7: "sedem",
    8: "osem",
    9: "deväť",
}

TEENS = {
    10: "desať",
    11: "jedenásť",
    12: "dvanásť",
    13: "trinásť",
    14: "štrnásť",
    15: "pätnásť",
    16: "šestnásť",
    17: "sedemnásť",
    18: "osemnásť",
    19: "devätnásť",
}

TENS = {
    20: "dvadsať",
    30: "tridsať",
    40: "štyridsať",
    50: "päťdesiat",
    60: "šesťdesiat",
    70: "sedemdesiat",
    80: "osemdesiat",
    90: "deväťdesiat",
}

HUNDREDS = {
    100: "sto",
    200: "dvesto",
    300: "tristo",
    400: "štyristo",
    500: "päťsto",
    600: "šesťsto",
    700: "sedemsto",
    800: "osemsto",
    900: "deväťsto",
}


def _under_100(n: int, one_word: str = "jeden", two_word: str = "dva") -> str:
    if n < 0 or n >= 100:
        raise ValueError("_under_100 expects 0..99")
    if n == 1:
        return one_word
    if n == 2:
        return two_word
    if n < 10:
        return ONES[n]
    if n < 20:
        return TEENS[n]
    tens = (n // 10) * 10
    rest = n % 10
    return TENS[tens] + (_under_100(rest) if rest else "")


def _under_1000(n: int) -> str:
    if n < 0 or n >= 1000:
        raise ValueError("_under_1000 expects 0..999")
    if n < 100:
        return _under_100(n)
    hundreds = (n // 100) * 100
    rest = n % 100
    return HUNDREDS[hundreds] + (_under_100(rest) if rest else "")


def number_to_slovak_words(n: int, *, one_word: str = "jeden", two_word: str = "dva") -> str:
    if n < 0 or n > 999_999:
        raise ValueError("Supported range is 0..999999")
    if n < 100:
        return _under_100(n, one_word=one_word, two_word=two_word)
    if n < 1000:
        return _under_1000(n)

    thousands = n // 1000
    rest = n % 1000

    if thousands == 1:
        prefix = "tisíc"
    elif thousands == 2:
        prefix = "dvetisíc"
    elif thousands == 3:
        prefix = "tritisíc"
    elif thousands == 4:
        prefix = "štyritisíc"
    else:
        prefix = _under_1000(thousands) + "tisíc"

    return prefix + (_under_1000(rest) if rest else "")


def euro_suffix(euros: int) -> str:
    if euros == 1:
        return "euro"
    last_two = euros % 100
    last = euros % 10
    if last in (2, 3, 4) and last_two not in (12, 13, 14):
        return "eurá"
    return "eur"


def cent_suffix(cents: int) -> str:
    if cents == 1:
        return "cent"
    last_two = cents % 100
    last = cents % 10
    if last in (2, 3, 4) and last_two not in (12, 13, 14):
        return "centy"
    return "centov"


def amount_in_words_sk(value: Decimal | float | int | str) -> str:
    value = money(value)
    euros = int(value)
    cents = int((value - Decimal(euros)) * 100)
    euro_words = number_to_slovak_words(euros, one_word="jedno", two_word="dve")
    cent_words = number_to_slovak_words(cents, one_word="jeden", two_word="dva")
    return f"{euro_words} {euro_suffix(euros)} a {cent_words} {cent_suffix(cents)}"


def date_sk(value: dt.date) -> str:
    return value.strftime("%d.%m.%Y")


def billing_month(value: dt.date) -> str:
    return value.strftime("%Y-%m")


def safe_filename(text: str) -> str:
    text = re.sub(r"[^\w\-.]+", "_", text, flags=re.UNICODE)
    text = re.sub(r"_+", "_", text).strip("._")
    return text[:120] or "invoice"


def compact_iban(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).upper()


def format_iban(value: str) -> str:
    compact = compact_iban(value)
    return " ".join(compact[i:i + 4] for i in range(0, len(compact), 4))


def mod97(number: str) -> int:
    remainder = 0
    for char in number:
        remainder = (remainder * 10 + int(char)) % 97
    return remainder


def iban_check_digits(country_code: str, bban: str) -> str:
    converted = "".join(str(ord(char) - 55) for char in country_code.upper())
    checksum = 98 - mod97(f"{bban}{converted}00")
    return f"{checksum:02d}"


def make_sk_iban(bank_code: str, account_number: str, prefix: str = "000000") -> str:
    bank_code = str(bank_code).zfill(4)
    prefix = str(prefix).zfill(6)
    account_number = str(account_number).zfill(10)
    bban = f"{bank_code}{prefix}{account_number}"
    check_digits = iban_check_digits("SK", bban)
    return format_iban(f"SK{check_digits}{bban}")


def parse_sk_iban(value: str) -> tuple[str, str] | None:
    compact = compact_iban(value)
    if not re.fullmatch(r"SK\d{22}", compact):
        return None
    bban = compact[4:]
    return bban[:4], bban[-10:]


def csv_json_list(values: Iterable[Decimal]) -> str:
    return json.dumps([decimal_string(value) for value in values], ensure_ascii=False)


# -----------------------------
# PDF helpers
# -----------------------------


PDF_FONT_CANDIDATES = [
    (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ),
    (
        Path("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf"),
    ),
]

# PDF palette — restrained navy accent on near-black ink, light panels.
PDF_INK = colors.HexColor("#16212E")
PDF_MUTED = colors.HexColor("#5B6B7C")
PDF_ACCENT = colors.HexColor("#1E3A5F")
PDF_ACCENT_TINT = colors.HexColor("#C9D6E5")
PDF_PANEL = colors.HexColor("#F2F5F9")
PDF_RULE = colors.HexColor("#D8DFE7")


@lru_cache(maxsize=1)
def get_pdf_fonts() -> tuple[str, str]:
    for regular_path, bold_path in PDF_FONT_CANDIDATES:
        if regular_path.exists() and bold_path.exists():
            regular_name = "InvoiceGenRegular"
            bold_name = "InvoiceGenBold"
            if regular_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
            if bold_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
            return regular_name, bold_name
    return "Helvetica", "Helvetica-Bold"


@lru_cache(maxsize=1)
def get_pdf_styles() -> dict[str, ParagraphStyle]:
    regular_font, bold_font = get_pdf_fonts()
    return {
        "body": ParagraphStyle(
            "body",
            fontName=regular_font,
            fontSize=9,
            leading=12,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "body_right": ParagraphStyle(
            "body_right",
            fontName=regular_font,
            fontSize=9,
            leading=12,
            alignment=TA_RIGHT,
            textColor=PDF_INK,
        ),
        "body_center": ParagraphStyle(
            "body_center",
            fontName=regular_font,
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
            textColor=PDF_INK,
        ),
        "small": ParagraphStyle(
            "small",
            fontName=regular_font,
            fontSize=8,
            leading=10.5,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "label": ParagraphStyle(
            "label",
            fontName=bold_font,
            fontSize=8,
            leading=10,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#334155"),
        ),
        "title": ParagraphStyle(
            "title",
            fontName=bold_font,
            fontSize=19,
            leading=23,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "heading": ParagraphStyle(
            "heading",
            fontName=bold_font,
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "heading_center": ParagraphStyle(
            "heading_center",
            fontName=bold_font,
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            textColor=PDF_INK,
        ),
        "value_emphasis": ParagraphStyle(
            "value_emphasis",
            fontName=bold_font,
            fontSize=11,
            leading=13,
            alignment=TA_RIGHT,
            textColor=PDF_INK,
        ),
        # --- redesigned-layout styles ---
        "doc_no_main": ParagraphStyle(
            "doc_no_main",
            fontName=bold_font,
            fontSize=12.5,
            leading=16,
            alignment=TA_RIGHT,
            textColor=colors.white,
        ),
        "doc_no_sub": ParagraphStyle(
            "doc_no_sub",
            fontName=regular_font,
            fontSize=9.5,
            leading=13,
            alignment=TA_RIGHT,
            textColor=PDF_ACCENT_TINT,
        ),
        "meta_label": ParagraphStyle(
            "meta_label",
            fontName=bold_font,
            fontSize=7.5,
            leading=10,
            alignment=TA_LEFT,
            textColor=PDF_MUTED,
        ),
        "meta_value": ParagraphStyle(
            "meta_value",
            fontName=bold_font,
            fontSize=9.5,
            leading=12,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "section_label": ParagraphStyle(
            "section_label",
            fontName=bold_font,
            fontSize=8.5,
            leading=11,
            alignment=TA_LEFT,
            textColor=PDF_MUTED,
        ),
        "party_name": ParagraphStyle(
            "party_name",
            fontName=bold_font,
            fontSize=10,
            leading=13,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "party_line": ParagraphStyle(
            "party_line",
            fontName=regular_font,
            fontSize=8.5,
            leading=11.5,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "th_left": ParagraphStyle(
            "th_left",
            fontName=bold_font,
            fontSize=8.5,
            leading=11,
            alignment=TA_LEFT,
            textColor=colors.white,
        ),
        "th_center": ParagraphStyle(
            "th_center",
            fontName=bold_font,
            fontSize=8.5,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.white,
        ),
        "th_right": ParagraphStyle(
            "th_right",
            fontName=bold_font,
            fontSize=8.5,
            leading=11,
            alignment=TA_RIGHT,
            textColor=colors.white,
        ),
        "total_label": ParagraphStyle(
            "total_label",
            fontName=bold_font,
            fontSize=9.5,
            leading=12,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
        "total_value": ParagraphStyle(
            "total_value",
            fontName=bold_font,
            fontSize=10,
            leading=13,
            alignment=TA_RIGHT,
            textColor=PDF_INK,
        ),
        "grand_label": ParagraphStyle(
            "grand_label",
            fontName=bold_font,
            fontSize=10,
            leading=13,
            alignment=TA_LEFT,
            textColor=colors.white,
        ),
        "grand_value": ParagraphStyle(
            "grand_value",
            fontName=bold_font,
            fontSize=12,
            leading=15,
            alignment=TA_RIGHT,
            textColor=colors.white,
        ),
        "words_line": ParagraphStyle(
            "words_line",
            fontName=regular_font,
            fontSize=8,
            leading=11,
            alignment=TA_RIGHT,
            textColor=PDF_MUTED,
        ),
        "note_text": ParagraphStyle(
            "note_text",
            fontName=regular_font,
            fontSize=8.5,
            leading=12,
            alignment=TA_LEFT,
            textColor=PDF_INK,
        ),
    }


def pdf_paragraph(text: str, style: str) -> Paragraph:
    styles = get_pdf_styles()
    safe = escape(str(text or "")).replace("\n", "<br/>")
    return Paragraph(safe, styles[style])


def pdf_markup(markup: str, style: str) -> Paragraph:
    styles = get_pdf_styles()
    return Paragraph(markup, styles[style])


def pdf_multiline(lines: list[str], style: str) -> Paragraph:
    escaped_lines = [escape(line) for line in lines if line]
    return pdf_markup("<br/>".join(escaped_lines), style)


def party_lines(party: Party, *, include_vat: bool) -> list[str]:
    lines = [
        party.name,
        party.street,
        party.city_country,
    ]
    if party.register:
        lines.append(party.register)
    if party.ico:
        lines.append(f"IČO: {party.ico}")
    if party.dic:
        lines.append(f"DIČ: {party.dic}")
    if party.ic_dph:
        lines.append(f"IČ DPH: {party.ic_dph}")
    if include_vat:
        lines.append(f"Platiteľ DPH: {party.vat_payer}")
    if party.contact:
        lines.append(party.contact)
    if party.phone:
        lines.append(f"Telefón: {party.phone}")
    if party.email:
        lines.append(f"Email: {party.email}")
    if party.web:
        lines.append(f"Web: {party.web}")
    if party.iban:
        lines.append(f"IBAN: {party.iban}")
    if party.swift:
        lines.append(f"SWIFT: {party.swift}")
    if party.bank_account:
        lines.append(f"Banka / účet: {party.bank_account}")
    return lines


def party_block(title: str, party: Party, *, include_vat: bool) -> Paragraph:
    lines = party_lines(party, include_vat=include_vat)
    markup_lines = [f"<b>{escape(title)}</b>", f"<b>{escape(lines[0])}</b>"]
    markup_lines.extend(escape(line) for line in lines[1:])
    return pdf_markup("<br/>".join(markup_lines), "body")


def party_groups(party: Party, *, include_vat: bool) -> list[list[str]]:
    # Same line texts as party_lines (minus the name), grouped so the PDF
    # can add breathing room between address / registry IDs / contact / bank.
    address = [party.street, party.city_country]
    if party.register:
        address.append(party.register)

    ids: list[str] = []
    if party.ico:
        ids.append(f"IČO: {party.ico}")
    if party.dic:
        ids.append(f"DIČ: {party.dic}")
    if party.ic_dph:
        ids.append(f"IČ DPH: {party.ic_dph}")
    if include_vat:
        ids.append(f"Platiteľ DPH: {party.vat_payer}")

    contact: list[str] = []
    if party.contact:
        contact.append(party.contact)
    if party.phone:
        contact.append(f"Telefón: {party.phone}")
    if party.email:
        contact.append(f"Email: {party.email}")
    if party.web:
        contact.append(f"Web: {party.web}")

    bank: list[str] = []
    if party.iban:
        bank.append(f"IBAN: {party.iban}")
    if party.swift:
        bank.append(f"SWIFT: {party.swift}")
    if party.bank_account:
        bank.append(f"Banka / účet: {party.bank_account}")

    return [group for group in (address, ids, contact, bank) if group]


def party_cell(party: Party, *, include_vat: bool) -> list[Any]:
    flowables: list[Any] = [pdf_paragraph(party.name, "party_name"), Spacer(1, 2 * mm)]
    groups = party_groups(party, include_vat=include_vat)
    for index, group in enumerate(groups):
        if index > 0:
            flowables.append(Spacer(1, 2 * mm))
        flowables.append(pdf_multiline(group, "party_line"))
    return flowables


def labeled_value(label: str, value: str) -> Paragraph:
    styles = get_pdf_styles()
    safe_label = escape(label)
    safe_value = escape(value)
    return Paragraph(f"<b>{safe_label}</b> {safe_value}", styles["body"])


def build_pdf_story(inv: InvoiceData, qr_png_path: Path) -> list[Any]:
    story: list[Any] = []

    # Header: big title left, navy document-number block right, accent rule.
    doc_no_cell = [
        pdf_paragraph(f"ZF / {inv.invoice_no}", "doc_no_main"),
        pdf_paragraph(f"VS: {inv.variable_symbol}", "doc_no_sub"),
    ]
    header = PdfTable(
        [[pdf_paragraph("ZÁLOHOVÁ FAKTÚRA", "title"), doc_no_cell]],
        colWidths=[112 * mm, 64 * mm],
    )
    header.setStyle(TableStyle([
        ("BACKGROUND", (1, 0), (1, 0), PDF_ACCENT),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (0, 0), 0),
        ("LEFTPADDING", (1, 0), (1, 0), 10),
        ("RIGHTPADDING", (1, 0), (1, 0), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(header)
    story.append(HRFlowable(width="100%", thickness=1.4, color=PDF_ACCENT, spaceBefore=3, spaceAfter=0))
    story.append(Spacer(1, 5 * mm))

    # Meta strip: label-over-value columns on a light panel, no grid.
    def meta_cell(label: str, value: str) -> list[Any]:
        return [pdf_paragraph(label, "meta_label"), pdf_paragraph(value, "meta_value")]

    meta = PdfTable(
        [[
            meta_cell("Forma úhrady:", "peňažný prevod"),
            meta_cell("Variabilný symbol:", inv.variable_symbol),
            meta_cell("Dátum vystavenia:", date_sk(inv.issue_date)),
            meta_cell("Dátum splatnosti:", date_sk(inv.due_date)),
        ]],
        colWidths=[44 * mm, 44 * mm, 44 * mm, 44 * mm],
    )
    meta.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), PDF_PANEL),
        ("LINEBEFORE", (1, 0), (-1, -1), 0.6, PDF_RULE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(meta)
    story.append(Spacer(1, 6 * mm))

    # Parties: two open columns with underlined section labels — no boxes.
    parties = PdfTable(
        [
            [pdf_paragraph("Dodávateľ", "section_label"), "", pdf_paragraph("Odberateľ", "section_label")],
            [party_cell(inv.supplier, include_vat=True), "", party_cell(inv.customer, include_vat=False)],
        ],
        colWidths=[84 * mm, 8 * mm, 84 * mm],
    )
    parties.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (0, 0), 0.9, PDF_ACCENT),
        ("LINEBELOW", (2, 0), (2, 0), 0.9, PDF_ACCENT),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, 0), 0),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 3),
        ("TOPPADDING", (0, 1), (-1, 1), 4),
        ("BOTTOMPADDING", (0, 1), (-1, 1), 0),
    ]))
    story.append(parties)
    story.append(Spacer(1, 7 * mm))

    # Items: dark header band, right-aligned numbers, horizontal rules only.
    items = PdfTable(
        [
            [
                pdf_paragraph("Popis položky", "th_left"),
                pdf_paragraph("Množstvo", "th_center"),
                pdf_paragraph("MJ", "th_center"),
                pdf_paragraph("Cena za MJ", "th_right"),
                pdf_paragraph("Celková cena", "th_right"),
            ],
            [
                pdf_paragraph(inv.item_description, "body"),
                pdf_paragraph(compact_decimal_string(inv.quantity), "body_center"),
                pdf_paragraph(inv.unit, "body_center"),
                pdf_paragraph(money_eur(inv.unit_price), "body_right"),
                pdf_paragraph(money_eur(inv.total), "body_right"),
            ],
        ],
        colWidths=[76 * mm, 24 * mm, 14 * mm, 30 * mm, 32 * mm],
        repeatRows=1,
    )
    items.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), PDF_ACCENT),
        ("LINEBELOW", (0, 1), (-1, 1), 0.6, PDF_RULE),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, 0), 5),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 5),
        ("TOPPADDING", (0, 1), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 7),
    ]))
    story.append(items)
    story.append(Spacer(1, 6 * mm))

    # Totals: compact right-aligned block, grand total on the accent band;
    # the amount in words gets a full-width line so it never wraps cramped.
    totals = PdfTable(
        [
            [pdf_paragraph("Spolu:", "total_label"), pdf_paragraph(money_eur(inv.total), "total_value")],
            [
                pdf_paragraph("K úhrade / zostáva uhradiť:", "grand_label"),
                pdf_paragraph(money_eur(inv.total), "grand_value"),
            ],
        ],
        colWidths=[66 * mm, 42 * mm],
        hAlign="RIGHT",
    )
    totals.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), PDF_PANEL),
        ("BACKGROUND", (0, 1), (-1, 1), PDF_ACCENT),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(totals)
    story.append(Spacer(1, 2 * mm))
    story.append(pdf_markup(
        f"<b>Slovom:</b> {escape(amount_in_words_sk(inv.total))}",
        "words_line",
    ))
    story.append(Spacer(1, 7 * mm))

    # Notes panel with a left accent bar; QR sits on the right.
    qr_image = PdfImage(str(qr_png_path), width=26 * mm, height=26 * mm)
    notes_cell = [
        pdf_paragraph("Poznámky / doplňujúci text", "section_label"),
        Spacer(1, 1.5 * mm),
        pdf_multiline(
            [
                inv.note,
                f"Referenčný text: {make_reference_text(inv.variable_symbol, billing_month(inv.issue_date))}",
            ],
            "note_text",
        ),
    ]
    notes_and_qr = PdfTable(
        [[notes_cell, qr_image]],
        colWidths=[140 * mm, 36 * mm],
    )
    notes_and_qr.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), PDF_PANEL),
        ("LINEBEFORE", (0, 0), (0, -1), 2.2, PDF_ACCENT),
        ("VALIGN", (0, 0), (0, 0), "TOP"),
        ("VALIGN", (1, 0), (1, 0), "MIDDLE"),
        ("ALIGN", (1, 0), (1, 0), "CENTER"),
        ("LEFTPADDING", (0, 0), (0, 0), 10),
        ("LEFTPADDING", (1, 0), (1, 0), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(notes_and_qr)
    story.append(Spacer(1, 16 * mm))

    # Signature area: two ruled columns with a gap, names beneath.
    footer = PdfTable(
        [
            [pdf_paragraph("Vyhotovil", "heading"), "", pdf_paragraph("Prevzal", "heading")],
            [pdf_paragraph(inv.prepared_by, "body"), "", pdf_paragraph(inv.received_by or "", "body")],
        ],
        colWidths=[80 * mm, 16 * mm, 80 * mm],
    )
    footer.setStyle(TableStyle([
        ("LINEABOVE", (0, 0), (0, 0), 0.8, PDF_MUTED),
        ("LINEABOVE", (2, 0), (2, 0), 0.8, PDF_MUTED),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, 0), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(footer)

    return story


def render_invoice_pdf(output_path: Path, inv: InvoiceData, qr_png_path: Path) -> None:
    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=17 * mm,
        rightMargin=17 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=f"Zalohova faktura {inv.invoice_no}",
        author="invoice_generator",
    )
    pdf.build(build_pdf_story(inv, qr_png_path))


# -----------------------------
# DOCX helpers
# -----------------------------


def iter_tables(container) -> Iterable[Table]:
    for table in container.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_tables(cell)


def table_text(table: Table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def find_table(doc: Document, *needles: str) -> Table:
    for table in iter_tables(doc):
        text = table_text(table)
        if all(needle in text for needle in needles):
            return table
    raise RuntimeError(f"Template section not found: {needles}. Did the template layout change?")


def clear_cell(cell) -> None:
    for child in list(cell._tc):
        if child.tag.endswith("tbl"):
            cell._tc.remove(child)
    cell.text = ""


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size_pt: Optional[float] = None,
    align: Optional[int] = None,
) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    if size_pt:
        run.font.size = Pt(size_pt)


def insert_qr(cell, png_path: Path, *, width_in: float = 1.35) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(png_path), width=Inches(width_in))


def make_qr_payload(inv: InvoiceData) -> str:
    return "\n".join([
        "ZALOHOVA FAKTURA",
        f"CISLO:{inv.invoice_no}",
        f"VS:{inv.variable_symbol}",
        f"SUMA:{money_number(inv.total)} EUR",
        f"IBAN:{inv.supplier.iban}",
        f"SWIFT:{inv.supplier.swift}",
        f"DODAVATEL:{inv.supplier.name}",
        f"ODBERATEL:{inv.customer.name}",
        f"SPLATNOST:{date_sk(inv.due_date)}",
    ])


def create_qr_png(inv: InvoiceData, out_path: Path) -> None:
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=8,
        border=2,
    )
    qr.add_data(make_qr_payload(inv))
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(out_path)


def fill_invoice_docx(template_path: Path, output_path: Path, inv: InvoiceData, qr_png_path: Path) -> None:
    doc = Document(str(template_path))

    header = find_table(doc, "ZÁLOHOVÁ FAKTÚRA", "ZF /")
    set_cell_text(header.rows[0].cells[1], f"ZF / {inv.invoice_no}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header.rows[0].cells[2], f"VS: {inv.variable_symbol}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    payment = find_table(doc, "Forma úhrady:", "Variabilný symbol:")
    set_cell_text(payment.rows[0].cells[3], inv.variable_symbol, align=WD_ALIGN_PARAGRAPH.CENTER)

    supplier = find_table(doc, "Dodávateľ", "Platiteľ DPH:", "Banka / účet:")
    set_cell_text(supplier.rows[1].cells[0], inv.supplier.name, bold=True)
    set_cell_text(supplier.rows[2].cells[0], inv.supplier.street)
    set_cell_text(supplier.rows[3].cells[0], inv.supplier.city_country)
    set_cell_text(supplier.rows[4].cells[0], inv.supplier.register)
    set_cell_text(supplier.rows[5].cells[1], inv.supplier.ico)
    set_cell_text(supplier.rows[6].cells[1], inv.supplier.dic)
    set_cell_text(supplier.rows[7].cells[1], inv.supplier.ic_dph)
    set_cell_text(supplier.rows[8].cells[1], inv.supplier.vat_payer)
    set_cell_text(supplier.rows[9].cells[1], inv.supplier.phone)
    set_cell_text(supplier.rows[10].cells[1], inv.supplier.email)
    set_cell_text(supplier.rows[11].cells[1], inv.supplier.web)
    set_cell_text(supplier.rows[12].cells[1], inv.supplier.iban)
    set_cell_text(supplier.rows[13].cells[1], inv.supplier.swift)
    set_cell_text(supplier.rows[14].cells[1], inv.supplier.bank_account)

    customer = find_table(doc, "Odberateľ", "kontaktná osoba")
    set_cell_text(customer.rows[1].cells[0], inv.customer.name, bold=True)
    set_cell_text(customer.rows[2].cells[0], inv.customer.street)
    set_cell_text(customer.rows[3].cells[0], inv.customer.city_country)
    set_cell_text(customer.rows[4].cells[1], inv.customer.ico)
    set_cell_text(customer.rows[5].cells[1], inv.customer.dic)
    set_cell_text(customer.rows[6].cells[1], inv.customer.ic_dph)
    set_cell_text(customer.rows[7].cells[0], inv.customer.register)
    set_cell_text(customer.rows[8].cells[0], inv.customer.contact)
    set_cell_text(customer.rows[9].cells[0], inv.customer.phone)
    set_cell_text(customer.rows[10].cells[0], inv.customer.email)
    set_cell_text(customer.rows[11].cells[0], inv.customer.note)

    dates = find_table(doc, "Dátum vystavenia", "Dátum splatnosti")
    set_cell_text(dates.rows[0].cells[1], date_sk(inv.issue_date), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(dates.rows[0].cells[3], date_sk(inv.due_date), align=WD_ALIGN_PARAGRAPH.CENTER)

    items = find_table(doc, "Popis položky", "Množstvo", "Cena za MJ")
    set_cell_text(items.rows[1].cells[0], inv.item_description)
    set_cell_text(items.rows[1].cells[1], compact_decimal_string(inv.quantity), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(items.rows[1].cells[2], inv.unit, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(items.rows[1].cells[3], money_eur(inv.unit_price), align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(items.rows[1].cells[4], money_eur(inv.total), align=WD_ALIGN_PARAGRAPH.RIGHT)

    for row_index in (2, 3):
        for column_index in range(5):
            set_cell_text(items.rows[row_index].cells[column_index], "")

    set_cell_text(items.rows[4].cells[4], money_eur(inv.total), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(items.rows[5].cells[4], money_eur(inv.total), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    notes = find_table(doc, "Poznámky / doplňujúci text")
    set_cell_text(notes.rows[0].cells[1], inv.note)

    qr_table = find_table(doc, "QR kód")
    insert_qr(qr_table.rows[0].cells[0], qr_png_path, width_in=1.35)

    footer = find_table(doc, "Vyhotovil", "Prevzal")
    set_cell_text(footer.rows[1].cells[0], inv.prepared_by)
    set_cell_text(footer.rows[1].cells[1], inv.received_by)

    for table in iter_tables(doc):
        text = table_text(table)
        if "[slovom]" in text or "K úhrade / zostáva uhradiť:" in text:
            if len(table.rows) == 2 and len(table.columns) == 2:
                set_cell_text(table.rows[0].cells[0], money_number(inv.total), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(table.rows[0].cells[1], "EUR", bold=True)
                set_cell_text(table.rows[1].cells[0], amount_in_words_sk(inv.total), size_pt=8)
                set_cell_text(table.rows[1].cells[1], "")

    doc.save(str(output_path))


# -----------------------------
# Sample/customer generation
# -----------------------------


SUPPLIER = Party(
    name="Demo Accounting, s.r.o.",
    street="Hlavná 12",
    city_country="821 01 Bratislava, Slovenská republika",
    register="Zapísaná v Obchodnom registri Mestského súdu Bratislava III, oddiel Sro, vložka 123456/B",
    ico="50123456",
    dic="2123456789",
    ic_dph="SK2123456789",
    vat_payer="Áno",
    phone="+421 900 123 456",
    email="fakturacia@demo-accounting.sk",
    web="www.demo-accounting.sk",
    iban="SK12 1100 0000 0029 8765 4321",
    swift="TATRSKBX",
    bank_account="Tatra banka / Demo Accounting, s.r.o.",
    bank_code="1100",
    account_number="2987654321",
    bank_name="Tatra banka",
)

ADJECTIVES = [
    "Modrá", "Zelená", "Tichá", "Rýchla", "Jasná", "Nová", "Mestská", "Dunajská", "Severná", "Slnečná",
    "Stredná", "Malá", "Veľká", "Digitálna", "Presná", "Horská", "Lúčná", "Moderná", "Pokojná", "Sivá",
]
NOUNS = [
    "Dielňa", "Kaviareň", "Logistika", "Servis", "Agentúra", "Obchod", "Ateliér", "Kancelária", "Výroba", "Poradňa",
    "Studio", "Projekt", "Centrum", "Sklad", "Technika", "Služby", "Marketing", "Architektúra", "Správa", "Vývoj",
]
CITIES = [
    ("Bratislava", "811 01"), ("Košice", "040 01"), ("Prešov", "080 01"), ("Žilina", "010 01"),
    ("Nitra", "949 01"), ("Trnava", "917 01"), ("Trenčín", "911 01"), ("Banská Bystrica", "974 01"),
    ("Komárno", "945 01"), ("Dunajská Streda", "929 01"),
]
STREETS = [
    "Hlavná", "Mlynská", "Školská", "Dlhá", "Krátka", "Kvetná", "Lesná", "Nová", "Poštová", "Obchodná",
]
FIRST_NAMES = ["Martin", "Peter", "Jana", "Lucia", "Zuzana", "Marek", "Tomáš", "Eva", "Andrea", "Michaela"]
LAST_NAMES = ["Novák", "Kováč", "Horváth", "Varga", "Tóth", "Nagy", "Baláž", "Molnár", "Farkaš", "Szabó"]
MONTHS_SK = [
    "január", "február", "marec", "apríl", "máj", "jún",
    "júl", "august", "september", "október", "november", "december",
]
BANK_DIRECTORY = [
    ("1100", "Tatra banka"),
    ("0200", "VUB banka"),
    ("0900", "Slovenska sporitelna"),
    ("8330", "Fio banka"),
    ("7500", "CSOB"),
]

AMOUNT_BUCKET_SPECS = [
    ("fixed_80", 300),
    ("fixed_180", 300),
    ("fixed_210", 250),
    ("random_whole_eur", 150),
]
PAYMENT_SCENARIO_SPECS = [
    ("exact_single", 400),
    ("exact_split_total", 200),
    ("underpay", 200),
    ("overpay", 200),
]
FIXED_BUCKET_AMOUNTS = {
    "fixed_80": Decimal("80.00"),
    "fixed_180": Decimal("180.00"),
    "fixed_210": Decimal("210.00"),
}
RANDOM_BUCKET_FORBIDDEN_AMOUNTS = {80, 180, 210}

ANNUAL_EXTRA_FEE = Decimal("120.00")
ANNUAL_EXTRA_INTERVAL_MONTHS = 12
ANNUAL_EXTRA_CUSTOMER_MODULO = 5
DIRECT_DEBIT_MODULO = 10
INACTIVE_STATUS_MODULO = 40
VAT_PAYER_TRUTHY_VALUES = {"áno", "ano", "yes", "true", "1"}


def first_non_empty(row: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = row.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


def allocate_counts(total: int, specs: list[tuple[str, int]]) -> dict[str, int]:
    weight_total = sum(weight for _, weight in specs)
    counts = {name: total * weight // weight_total for name, weight in specs}
    remainder = total - sum(counts.values())
    ranked = sorted(
        (
            ((total * weight / weight_total) - counts[name], order, name)
            for order, (name, weight) in enumerate(specs)
        ),
        key=lambda item: (-item[0], item[1]),
    )
    for _, _, name in ranked[:remainder]:
        counts[name] += 1
    return counts


def build_weighted_sequence(total: int, specs: list[tuple[str, int]]) -> list[str]:
    counts = allocate_counts(total, specs)
    used = {name: 0 for name, _ in specs}
    sequence: list[str] = []

    for position in range(1, total + 1):
        chosen_name = ""
        chosen_score = float("-inf")
        chosen_order = len(specs)

        for order, (name, _) in enumerate(specs):
            if used[name] >= counts[name]:
                continue
            score = counts[name] * position / total - used[name]
            if score > chosen_score or (score == chosen_score and chosen_name == ""):
                chosen_name = name
                chosen_score = score
                chosen_order = order
            elif score == chosen_score and order < chosen_order:
                chosen_name = name
                chosen_order = order

        if not chosen_name:
            raise RuntimeError("Failed to build weighted sequence")

        used[chosen_name] += 1
        sequence.append(chosen_name)

    return sequence


def generate_bank_identity(index: int, seed: int) -> tuple[str, str, str, str]:
    bank_code, bank_name = BANK_DIRECTORY[(index + abs(seed)) % len(BANK_DIRECTORY)]
    account_number = f"{7000000000 + abs(seed) * 10000 + index:010d}"[-10:]
    iban = make_sk_iban(bank_code, account_number)
    return bank_code, account_number, bank_name, iban


def generated_customer(index: int, seed: int) -> Party:
    seed_offset = abs(seed)
    city, zip_code = CITIES[(index + seed_offset) % len(CITIES)]
    adjective = ADJECTIVES[(index + seed_offset) % len(ADJECTIVES)]
    noun = NOUNS[((index // len(ADJECTIVES)) + seed_offset) % len(NOUNS)]
    street = STREETS[(index * 3 + seed_offset) % len(STREETS)]
    first = FIRST_NAMES[(index * 7 + seed_offset) % len(FIRST_NAMES)]
    last = LAST_NAMES[(index * 11 + seed_offset) % len(LAST_NAMES)]
    bank_code, account_number, bank_name, iban = generate_bank_identity(index, seed)

    ico = f"{60000000 + index:08d}"
    dic = f"212{index:07d}"[:10]
    ic_dph = "" if index % 3 else f"SK{dic}"
    customer_id = f"CUST-{index:05d}"
    street_number = 100 + index

    return Party(
        name=f"{adjective} {noun} {index:04d}, s.r.o.",
        street=f"{street} {street_number}",
        city_country=f"{zip_code} {city}, Slovenská republika",
        ico=ico,
        dic=dic,
        ic_dph=ic_dph,
        vat_payer="Áno" if ic_dph else "Nie",
        register=f"Zapísaná v Obchodnom registri príslušného súdu, oddiel Sro, vložka {10000 + index}/X",
        contact=f"Kontaktná osoba: {first} {last}",
        phone=f"+421 9{index % 10:01d}{(1000000 + index) % 9000000:07d}",
        email=f"fakturacia+{index:05d}@example.test",
        note="Vzorový zákazník pre testovanie dávkového generovania.",
        customer_id=customer_id,
        iban=iban,
        bank_code=bank_code,
        account_number=account_number,
        bank_name=bank_name,
        bank_account=f"{bank_name} / {customer_id}",
    )


def normalize_customer_row(row: dict[str, str], index: int, seed: int) -> Party:
    synthetic = generated_customer(index, seed)
    provided_iban = first_non_empty(row, "customer_iban", "iban")
    provided_bank_code = first_non_empty(row, "customer_bank_code", "bank_code")
    provided_account_number = first_non_empty(row, "customer_account_number", "account_number")
    provided_bank_name = first_non_empty(row, "customer_bank_name", "bank_name")
    parsed = parse_sk_iban(provided_iban) if provided_iban else None

    bank_code = provided_bank_code or (parsed[0] if parsed else synthetic.bank_code)
    account_number = provided_account_number or (parsed[1] if parsed else synthetic.account_number)
    bank_name = provided_bank_name
    if not bank_name:
        bank_name = next((name for code, name in BANK_DIRECTORY if code == bank_code), synthetic.bank_name)

    iban = format_iban(provided_iban) if provided_iban else make_sk_iban(bank_code, account_number)

    return Party(
        name=first_non_empty(row, "customer_name", "name") or synthetic.name,
        street=first_non_empty(row, "street") or synthetic.street,
        city_country=first_non_empty(row, "city_country") or synthetic.city_country,
        ico=first_non_empty(row, "ico") or synthetic.ico,
        dic=first_non_empty(row, "dic") or synthetic.dic,
        ic_dph=first_non_empty(row, "ic_dph") or synthetic.ic_dph,
        register=first_non_empty(row, "register") or synthetic.register,
        vat_payer=first_non_empty(row, "vat_payer") or synthetic.vat_payer,
        phone=first_non_empty(row, "phone") or synthetic.phone,
        email=first_non_empty(row, "email") or synthetic.email,
        web=first_non_empty(row, "web") or synthetic.web,
        iban=iban,
        swift=first_non_empty(row, "swift") or synthetic.swift,
        bank_account=first_non_empty(row, "bank_account") or f"{bank_name} / {(first_non_empty(row, 'customer_id') or synthetic.customer_id)}",
        contact=first_non_empty(row, "contact") or synthetic.contact,
        note=first_non_empty(row, "note") or synthetic.note,
        customer_id=first_non_empty(row, "customer_id", "id") or synthetic.customer_id,
        bank_code=bank_code,
        account_number=account_number,
        bank_name=bank_name,
        annual_extra_fee=first_non_empty(row, "annual_extra_fee"),
        annual_extra_interval_months=first_non_empty(row, "annual_extra_interval_months"),
        payment_method=first_non_empty(row, "payment_method"),
        vat_status=first_non_empty(row, "vat_status"),
        status=first_non_empty(row, "status"),
    )


def read_customer_rows(path: Path) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows.append({key: str(value or "").strip() for key, value in row.items()})
    if not rows:
        raise ValueError(f"No customers found in {path}")
    return rows


def resolve_customers(count: int, start_index: int, seed: int, customers_csv: Optional[Path]) -> list[Party]:
    csv_rows = read_customer_rows(customers_csv) if customers_csv else []
    customers: list[Party] = []

    for offset in range(count):
        index = start_index + offset
        if offset < len(csv_rows):
            customers.append(normalize_customer_row(csv_rows[offset], index, seed))
        else:
            customers.append(generated_customer(index, seed))

    return customers


def random_bucket_amount(rng: random.Random, *, minimum: int = 1) -> Decimal:
    while True:
        candidate = rng.randint(minimum, 9999)
        if candidate not in RANDOM_BUCKET_FORBIDDEN_AMOUNTS:
            return money(candidate)


def payment_offset(invoice_total: Decimal) -> Decimal:
    euros = int(invoice_total)
    if euros <= 99:
        return Decimal("10.00")
    if euros <= 999:
        return Decimal("20.00")
    return Decimal("50.00")


def exact_split_amounts(invoice_total: Decimal, index: int) -> tuple[Decimal, Decimal]:
    euros = int(invoice_total)
    if euros < 2:
        raise ValueError("Split payments require an amount of at least 2 EUR")
    adjustment = (index % 5) - 2
    first = max(1, min(euros - 1, euros // 2 + adjustment))
    second = euros - first
    return money(first), money(second)


def planned_payment(invoice_total: Decimal, scenario: str, index: int) -> tuple[Decimal, tuple[Decimal, ...]]:
    if scenario == "exact_single":
        return invoice_total, ()
    if scenario == "exact_split_total":
        return invoice_total, exact_split_amounts(invoice_total, index)
    if scenario == "underpay":
        received = max(Decimal("1.00"), invoice_total - payment_offset(invoice_total))
        if received >= invoice_total:
            received = money(invoice_total - Decimal("1.00"))
        return received, ()
    if scenario == "overpay":
        return invoice_total + payment_offset(invoice_total), ()
    raise ValueError(f"Unsupported payment scenario: {scenario}")


def has_annual_extra(index: int) -> bool:
    return index % ANNUAL_EXTRA_CUSTOMER_MODULO == 0


def customer_payment_method(index: int) -> str:
    return "direct_debit" if index % DIRECT_DEBIT_MODULO == 0 else "bank_transfer"


def customer_status(index: int) -> str:
    return "inactive" if index % INACTIVE_STATUS_MODULO == 0 else "active"


def vat_status_for(party: Party) -> str:
    if party.ic_dph.strip() or party.vat_payer.strip().lower() in VAT_PAYER_TRUTHY_VALUES:
        return "vat_payer"
    return "non_vat_payer"


def billing_profile_party(customer: Party, index: int, amount: Decimal) -> Party:
    annual_extra = has_annual_extra(index)
    return replace(
        customer,
        monthly_fee=decimal_string(amount),
        annual_extra_fee=customer.annual_extra_fee or (decimal_string(ANNUAL_EXTRA_FEE) if annual_extra else ""),
        annual_extra_interval_months=customer.annual_extra_interval_months
        or (str(ANNUAL_EXTRA_INTERVAL_MONTHS) if annual_extra else ""),
        payment_method=customer.payment_method or customer_payment_method(index),
        vat_status=customer.vat_status or vat_status_for(customer),
        status=customer.status or customer_status(index),
    )


def amount_for_bucket(bucket: str, scenario: str, rng: random.Random) -> Decimal:
    if bucket in FIXED_BUCKET_AMOUNTS:
        return FIXED_BUCKET_AMOUNTS[bucket]
    minimum = 2 if scenario in {"exact_split_total", "underpay"} else 1
    return random_bucket_amount(rng, minimum=minimum)


def make_reference_text(variable_symbol: str, month: str) -> str:
    return f"/VS{variable_symbol}/KS0308/TXT {month} accounting services"


def build_invoice(index: int, customer: Party, issue_date: dt.date, amount: Decimal) -> InvoiceData:
    variable_symbol = f"{issue_date.year}{index:04d}"
    month_name = MONTHS_SK[(issue_date.month - 1) % 12]
    return InvoiceData(
        invoice_no=f"{issue_date.year}-{index:04d}",
        variable_symbol=variable_symbol,
        issue_date=issue_date,
        due_date=issue_date + dt.timedelta(days=14),
        supplier=SUPPLIER,
        customer=customer,
        item_description=f"Účtovnícke práce za {month_name} {issue_date.year}",
        quantity=Decimal("1"),
        unit="ks",
        unit_price=amount,
        prepared_by="Jana Účtovníková",
        received_by="",
    )


def build_annual_extra_invoice(index: int, customer: Party, issue_date: dt.date, amount: Decimal) -> InvoiceData:
    # 9-digit VS (year + "9" + index) cannot collide with the 8-digit monthly VS scheme.
    variable_symbol = f"{issue_date.year}9{index:04d}"
    return InvoiceData(
        invoice_no=f"{issue_date.year}-{index:04d}-AE",
        variable_symbol=variable_symbol,
        issue_date=issue_date,
        due_date=issue_date + dt.timedelta(days=14),
        supplier=SUPPLIER,
        customer=customer,
        item_description=f"Ročný doplnkový poplatok za {issue_date.year}",
        quantity=Decimal("1"),
        unit="ks",
        unit_price=amount,
        prepared_by="Jana Účtovníková",
        received_by="",
    )


def build_batch_id(issue_date: dt.date, start_index: int, count: int, seed: int) -> str:
    return f"invgen-{issue_date:%Y%m%d}-start{start_index:05d}-count{count:05d}-seed{seed}"


def build_batch_plan(
    *,
    count: int,
    start_index: int,
    issue_date: dt.date,
    seed: int,
    customers_csv: Optional[Path],
) -> list[PlannedInvoice]:
    rng = random.Random(seed)
    customers = resolve_customers(count, start_index, seed, customers_csv)
    amount_buckets = build_weighted_sequence(count, AMOUNT_BUCKET_SPECS)
    payment_scenarios = build_weighted_sequence(count, PAYMENT_SCENARIO_SPECS)
    batch_id = build_batch_id(issue_date, start_index, count, seed)
    records: list[PlannedInvoice] = []

    annual_extra_records: list[PlannedInvoice] = []

    for offset, customer in enumerate(customers):
        index = start_index + offset
        amount_bucket = amount_buckets[offset]
        payment_scenario = payment_scenarios[offset]
        amount = amount_for_bucket(amount_bucket, payment_scenario, rng)
        customer = billing_profile_party(customer, index, amount)
        invoice = build_invoice(index, customer, issue_date, amount)
        simulated_paid_total, simulated_split_amounts = planned_payment(invoice.total, payment_scenario, index)
        filename_base = safe_filename(
            f"ZF_{invoice.invoice_no}_{invoice.variable_symbol}_{customer.customer_id}_{customer.name}"
        )
        records.append(PlannedInvoice(
            batch_id=batch_id,
            invoice=invoice,
            charge_type="monthly",
            amount_bucket=amount_bucket,
            payment_scenario=payment_scenario,
            simulated_paid_total=simulated_paid_total,
            simulated_split_amounts=simulated_split_amounts,
            reference_text_template=make_reference_text(invoice.variable_symbol, billing_month(issue_date)),
            filename_base=filename_base,
        ))

        if has_annual_extra(index):
            extra_amount = money(customer.annual_extra_fee or ANNUAL_EXTRA_FEE)
            extra_invoice = build_annual_extra_invoice(index, customer, issue_date, extra_amount)
            extra_filename_base = safe_filename(
                f"ZF_{extra_invoice.invoice_no}_{extra_invoice.variable_symbol}_{customer.customer_id}_{customer.name}"
            )
            annual_extra_records.append(PlannedInvoice(
                batch_id=batch_id,
                invoice=extra_invoice,
                charge_type="annual_extra",
                amount_bucket="annual_extra",
                payment_scenario="exact_single",
                simulated_paid_total=extra_invoice.total,
                simulated_split_amounts=(),
                reference_text_template=make_reference_text(extra_invoice.variable_symbol, billing_month(issue_date)),
                filename_base=extra_filename_base,
            ))

    # Appended after the main loop so the monthly batch stays byte-identical per seed.
    records.extend(annual_extra_records)
    return records


# -----------------------------
# Manifest writing
# -----------------------------


def supplier_manifest() -> dict[str, str]:
    return {
        "supplier_name": SUPPLIER.name,
        "supplier_ico": SUPPLIER.ico,
        "supplier_dic": SUPPLIER.dic,
        "supplier_ic_dph": SUPPLIER.ic_dph,
        "supplier_iban": SUPPLIER.iban_compact,
        "supplier_swift": SUPPLIER.swift,
        "supplier_bank_code": SUPPLIER.bank_code,
        "supplier_account_number": SUPPLIER.account_number,
        "supplier_bank_name": SUPPLIER.bank_name,
        "supplier_email": SUPPLIER.email,
    }


def customer_manifest_row(customer: Party) -> dict[str, str]:
    return {
        "customer_id": customer.customer_id,
        "customer_name": customer.name,
        "street": customer.street,
        "city_country": customer.city_country,
        "ico": customer.ico,
        "dic": customer.dic,
        "ic_dph": customer.ic_dph,
        "register": customer.register,
        "contact": customer.contact,
        "phone": customer.phone,
        "email": customer.email,
        "customer_iban": customer.iban_compact,
        "customer_bank_code": customer.bank_code,
        "customer_account_number": customer.account_number,
        "customer_bank_name": customer.bank_name,
        "monthly_fee": customer.monthly_fee,
        "annual_extra_fee": customer.annual_extra_fee,
        "annual_extra_interval_months": customer.annual_extra_interval_months,
        "payment_method": customer.payment_method,
        "vat_status": customer.vat_status,
        "status": customer.status,
    }


def expected_charge_manifest_row(record: PlannedInvoice) -> dict[str, str]:
    invoice = record.invoice
    customer = invoice.customer
    return {
        "batch_id": record.batch_id,
        "customer_id": customer.customer_id,
        "customer_name": customer.name,
        "billing_month": billing_month(invoice.issue_date),
        "charge_type": record.charge_type,
        "variable_symbol": invoice.variable_symbol,
        "charge_amount": decimal_string(invoice.total),
        "due_date": invoice.due_date.isoformat(),
        "invoice_no": invoice.invoice_no,
    }


def invoice_manifest_row(record: PlannedInvoice, *, include_pdf: bool) -> dict[str, str]:
    invoice = record.invoice
    customer = invoice.customer
    return {
        "batch_id": record.batch_id,
        "customer_id": customer.customer_id,
        "customer_name": customer.name,
        "variable_symbol": invoice.variable_symbol,
        "billing_month": billing_month(invoice.issue_date),
        "issue_date": invoice.issue_date.isoformat(),
        "due_date": invoice.due_date.isoformat(),
        "invoice_total_amount": decimal_string(invoice.total),
        "currency": "EUR",
        "invoice_no": invoice.invoice_no,
        "charge_type": record.charge_type,
        "item_description": invoice.item_description,
        "quantity": compact_decimal_string(invoice.quantity),
        "unit": invoice.unit,
        "unit_price": decimal_string(invoice.unit_price),
        "amount_bucket": record.amount_bucket,
        "payment_scenario": record.payment_scenario,
        "simulated_paid_total": decimal_string(record.simulated_paid_total),
        "simulated_split_amounts": csv_json_list(record.simulated_split_amounts),
        "supplier_name": invoice.supplier.name,
        "supplier_iban": invoice.supplier.iban_compact,
        "customer_iban": customer.iban_compact,
        "customer_bank_code": customer.bank_code,
        "customer_account_number": customer.account_number,
        "street": customer.street,
        "city_country": customer.city_country,
        "ico": customer.ico,
        "dic": customer.dic,
        "ic_dph": customer.ic_dph,
        "register": customer.register,
        "contact": customer.contact,
        "phone": customer.phone,
        "email": customer.email,
        "docx_filename": record.docx_filename,
        "pdf_filename": record.pdf_filename if include_pdf else "",
        "reference_text_template": record.reference_text_template,
    }


def json_invoice_record(record: PlannedInvoice, *, include_pdf: bool) -> dict[str, Any]:
    row = invoice_manifest_row(record, include_pdf=include_pdf)
    row["simulated_split_amounts"] = [decimal_string(value) for value in record.simulated_split_amounts]
    return row


def write_csv(path: Path, rows: list[dict[str, str]]) -> None:
    if not rows:
        raise ValueError(f"No rows to write to {path}")
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_manifests(outdir: Path, records: list[PlannedInvoice], *, include_pdf: bool, template: Path, seed: int) -> None:
    manifests_dir = outdir / "manifests"
    manifests_dir.mkdir(parents=True, exist_ok=True)

    invoice_rows = [invoice_manifest_row(record, include_pdf=include_pdf) for record in records]
    expected_charge_rows = [expected_charge_manifest_row(record) for record in records]
    customers_by_id = {record.invoice.customer.customer_id: record.invoice.customer for record in records}
    customer_rows = [customer_manifest_row(customer) for customer in customers_by_id.values()]

    batch = {
        "batch_id": records[0].batch_id if records else "",
        "count": len(records),
        "issue_date": records[0].invoice.issue_date.isoformat() if records else "",
        "billing_month": billing_month(records[0].invoice.issue_date) if records else "",
        "template": str(template),
        "seed": seed,
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
    }

    write_csv(manifests_dir / "invoices.csv", invoice_rows)
    write_csv(manifests_dir / "expected_charges.csv", expected_charge_rows)
    write_csv(manifests_dir / "customers.csv", customer_rows)

    invoices_json = {
        "batch": batch,
        "supplier": supplier_manifest(),
        "invoices": [json_invoice_record(record, include_pdf=include_pdf) for record in records],
    }
    expected_charges_json = {
        "batch": batch,
        "expected_charges": expected_charge_rows,
    }
    customers_json = {
        "batch": batch,
        "customers": customer_rows,
    }

    with (manifests_dir / "invoices.json").open("w", encoding="utf-8") as handle:
        json.dump(invoices_json, handle, ensure_ascii=False, indent=2)
    with (manifests_dir / "expected_charges.json").open("w", encoding="utf-8") as handle:
        json.dump(expected_charges_json, handle, ensure_ascii=False, indent=2)
    with (manifests_dir / "customers.json").open("w", encoding="utf-8") as handle:
        json.dump(customers_json, handle, ensure_ascii=False, indent=2)


# -----------------------------
# CLI
# -----------------------------


def parse_date(value: str) -> dt.date:
    try:
        return dt.datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError as exc:
        raise argparse.ArgumentTypeError("Use YYYY-MM-DD, e.g. 2026-05-25") from exc


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate many filled DOCX/PDF advance invoice documents.")
    parser.add_argument("--template", required=True, type=Path, help="Path to the DOCX template.")
    parser.add_argument("--outdir", default=Path("generated_invoices"), type=Path, help="Output directory.")
    parser.add_argument("--count", default=1000, type=int, help="Number of documents to generate.")
    parser.add_argument("--start-index", default=1, type=int, help="First invoice/customer index.")
    parser.add_argument("--issue-date", default=dt.date.today(), type=parse_date, help="Issue date YYYY-MM-DD.")
    parser.add_argument("--customers-csv", type=Path, help="Optional CSV with customer data. Missing rows are backfilled.")
    parser.add_argument("--pdf", action="store_true", help="Also export text-based PDFs directly from Python.")
    parser.add_argument("--soffice", help="Deprecated compatibility flag. Ignored by the direct PDF renderer.")
    parser.add_argument("--seed", default=42, type=int, help="Random seed for deterministic demo data.")
    args = parser.parse_args()

    if args.count <= 0:
        raise SystemExit("--count must be greater than zero")
    if not args.template.exists():
        raise SystemExit(f"Template not found: {args.template}")

    outdir: Path = args.outdir
    docx_dir = outdir / "docx"
    pdf_dir = outdir / "pdf"
    qr_dir = outdir / "_qr_tmp"
    docx_dir.mkdir(parents=True, exist_ok=True)
    qr_dir.mkdir(parents=True, exist_ok=True)

    records = build_batch_plan(
        count=args.count,
        start_index=args.start_index,
        issue_date=args.issue_date,
        seed=args.seed,
        customers_csv=args.customers_csv,
    )

    generated_docx: list[Path] = []
    generated_pdf: list[Path] = []

    for index, record in enumerate(records, start=1):
        qr_path = qr_dir / f"{record.filename_base}.png"
        docx_path = docx_dir / record.docx_filename
        pdf_path = pdf_dir / record.pdf_filename

        create_qr_png(record.invoice, qr_path)
        fill_invoice_docx(args.template, docx_path, record.invoice, qr_path)
        generated_docx.append(docx_path)
        if args.pdf:
            pdf_dir.mkdir(parents=True, exist_ok=True)
            render_invoice_pdf(pdf_path, record.invoice, qr_path)
            generated_pdf.append(pdf_path)

        if index % 50 == 0:
            print(f"Generated {index} DOCX files...", flush=True)

    print(f"Generated {len(generated_docx)} DOCX files in: {docx_dir}")

    if args.pdf:
        print(f"Generated {len(generated_pdf)} PDFs in: {pdf_dir}")

    write_manifests(outdir, records, include_pdf=args.pdf, template=args.template, seed=args.seed)
    print(f"Generated manifests in: {outdir / 'manifests'}")
    print("Done. Bureaucracy successfully multiplied, because apparently that was the mission.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
